"""Format-preserving DOCX resume tailoring.

Keeps the original .docx's exact structure (paragraph order/count, styles,
fonts, bold, spacing) - the LLM only proposes new wording for content
paragraphs (the summary, achievement/responsibility bullets). python-docx
never touches formatting; it only swaps the text inside existing runs, so
fonts/sizes/bold can never drift from the original.

Real documents are messier than their style names suggest - in the seed
resume, the same paragraph style is reused for both section headers and
bullets in places, so a style-name heuristic can't reliably tell "structural"
from "content" apart. The LLM classifies each paragraph instead (a judgment
call it's suited for); the mechanical text substitution stays deterministic.

Page length is enforced via a real, checkable proxy rather than a guess:
Word caches the last-known page/word/character count in docProps/app.xml
every time it saves a .docx. Nothing is added or removed here - only
reworded - so keeping each rewrite at or under its original paragraph's
character count is what keeps the tailored version at the same page count.
"""

from __future__ import annotations

import io
import os
import re
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Literal, Optional

import docx
import httpx
from google import genai
from google.genai import errors as genai_errors
from pydantic import BaseModel, ValidationError

MODEL = "gemini-3.5-flash"


def _atomic_write_bytes(path: Path, data: bytes) -> None:
    """Write `data` to `path` atomically - mirrors `jobs.cli._atomic_write_text`'s
    design (same-directory pid+uuid tmp naming, fsync, `os.replace`, cleanup on
    failure) but for binary content, since `Document.save()` writes a zip
    container, not text. Implemented locally rather than imported from
    `jobs.cli`, which already imports this module - importing back would be
    circular."""
    tmp = path.parent / f"{path.name}.{os.getpid()}.{uuid.uuid4().hex}.tmp"
    try:
        with open(tmp, "wb") as f:
            f.write(data)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp, path)
    except OSError:
        tmp.unlink(missing_ok=True)
        raise


@dataclass(frozen=True)
class SourceParagraph:
    index: int
    text: str


def extract_paragraphs(docx_path: str | Path) -> list[SourceParagraph]:
    document = docx.Document(str(docx_path))
    return [SourceParagraph(index=i, text=p.text) for i, p in enumerate(document.paragraphs) if p.text.strip()]


def get_cached_document_stats(docx_path: str | Path) -> Optional[dict]:
    """Word's own last-saved page/word/character counts, read straight from
    the docx's docProps/app.xml - a real signal, not an estimate. Returns
    None if the file was never saved by an app that writes these (rare)."""
    try:
        with zipfile.ZipFile(str(docx_path)) as archive:
            xml = archive.read("docProps/app.xml").decode("utf-8")
    except (KeyError, OSError):
        return None

    stats = {}
    for key in ("Pages", "Words", "Characters", "CharactersWithSpaces"):
        match = re.search(rf"<{key}>(\d+)</{key}>", xml)
        if match:
            stats[key.lower()] = int(match.group(1))
    return stats or None


class ParagraphEdit(BaseModel):
    index: int
    action: Literal["keep", "rewrite"]
    text: Optional[str] = None


class TailoredResumeEdits(BaseModel):
    edits: list[ParagraphEdit]


_SYSTEM_INSTRUCTION = (
    "You tailor a candidate's resume to a specific job posting by rewording "
    "SOME paragraphs, while others must stay byte-for-byte unchanged. You are "
    "given every non-empty paragraph in the original .docx, numbered by "
    "index, with its original character length. For each paragraph, decide:\n"
    "- 'keep': structural lines that must NEVER be touched - the candidate's "
    "name, contact info/tagline, section headings (SUMMARY, TECHNICAL "
    "SKILLS, WORK EXPERIENCE, EDUCATION, etc.), job title/company/date "
    "header lines, project title lines, education and certification "
    "entries, GitHub/URL lines.\n"
    "- 'rewrite': content lines where wording can be tailored - the summary "
    "paragraph, and achievement/responsibility bullets under a role or "
    "project.\n"
    "Rules for every 'rewrite': keep the same real facts and metrics (never "
    "invent a number, technology, or achievement not already present "
    "somewhere in the resume); emphasize what's relevant to this job; and "
    "critically, keep the new text AT OR UNDER the original paragraph's "
    "character count - the document layout is fixed and must stay within "
    "its original page count, so a longer replacement will overflow the "
    "page. When in doubt, prefer 'keep' - don't reword something you're not "
    "confident is content rather than structure. Return an edit for every "
    "paragraph index given, even the ones you keep."
)


def _build_input(paragraphs: list[SourceParagraph], job_raw_text: str, company_name: Optional[str]) -> str:
    lines = [f"[{p.index}] ({len(p.text)} chars) {p.text}" for p in paragraphs]
    return (
        "ORIGINAL RESUME PARAGRAPHS:\n"
        + "\n".join(lines)
        + f"\n\nTARGET COMPANY: {company_name or 'unknown'}\n\nJOB POSTING:\n{job_raw_text}\n"
    )


def generate_paragraph_edits(
    paragraphs: list[SourceParagraph],
    job_raw_text: str,
    company_name: Optional[str],
    *,
    client: Optional[genai.Client] = None,
) -> dict[int, str]:
    """Returns {paragraph_index: new_text} only for paragraphs marked 'rewrite'."""
    client = client or genai.Client()
    try:
        interaction = client.interactions.create(
            model=MODEL,
            system_instruction=_SYSTEM_INSTRUCTION,
            input=_build_input(paragraphs, job_raw_text, company_name),
            response_format={
                "type": "text",
                "mime_type": "application/json",
                "schema": TailoredResumeEdits.model_json_schema(),
            },
        )
        result = TailoredResumeEdits.model_validate_json(interaction.output_text)
        return {edit.index: edit.text for edit in result.edits if edit.action == "rewrite" and edit.text}
    except (
        genai_errors.APIError,
        genai_errors.UnknownApiResponseError,
        httpx.HTTPError,
        ValidationError,
        RuntimeError,  # covers the SDK's bare RuntimeError when no API credentials resolve
    ) as exc:
        detail = str(exc).strip() or type(exc).__name__
        raise SystemExit(f"Resume paragraph tailoring failed: {detail}") from exc


def _set_paragraph_text_preserving_format(paragraph, new_text: str) -> None:
    """Replace a paragraph's visible text while keeping its exact formatting.

    Reuses the first run's properties (font, size, bold, color) for the new
    text rather than creating a fresh run with default styling, and blanks
    any extra runs. Known limitation: a rewritten paragraph collapses to one
    run's formatting, so mid-paragraph emphasis (e.g. a bolded product name
    inside an otherwise plain bullet) isn't reapplied to the new wording.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def build_tailored_docx(source_path: str | Path, rewritten: dict[int, str], out_path: str | Path) -> None:
    """Apply rewritten paragraph text onto a fresh copy of the source .docx -
    the original file on disk is never modified."""
    document = docx.Document(str(source_path))
    for i, paragraph in enumerate(document.paragraphs):
        if i in rewritten:
            _set_paragraph_text_preserving_format(paragraph, rewritten[i])
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    buffer = io.BytesIO()
    document.save(buffer)
    _atomic_write_bytes(out_path, buffer.getvalue())


def estimate_page_risk(source_path: str | Path, rewritten: dict[int, str]) -> Optional[str]:
    """A grounded (not guessed) length-overflow check: Word's own cached
    chars-per-page ratio from the original, applied to the new total. Can't
    render the .docx to check pages directly, so this says what it can
    verify and stays quiet rather than claiming false precision."""
    stats = get_cached_document_stats(source_path)
    if not stats or not stats.get("pages") or not stats.get("characterswithspaces"):
        return None

    original_chars = stats["characterswithspaces"]
    pages = stats["pages"]

    paragraphs = extract_paragraphs(source_path)
    new_total = sum(len(rewritten.get(p.index, p.text)) for p in paragraphs)
    original_total = sum(len(p.text) for p in paragraphs)
    delta = new_total - original_total

    projected_chars = original_chars + delta
    chars_per_page = original_chars / pages
    projected_pages = projected_chars / chars_per_page

    if projected_pages > pages + 0.15:
        return (
            f"Tailored resume text is {delta:+d} chars vs the original "
            f"({original_total} -> {new_total}) - may push past the original "
            f"{pages}-page length. Open it and check before sending."
        )
    return None


def write_plain_docx(text: str, out_path: str | Path) -> None:
    """A simple, cleanly-formatted .docx for content with no source template
    to preserve (the cover letter) - one paragraph per blank-line-separated
    block, using python-docx's default document style throughout."""
    document = docx.Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            document.add_paragraph(block)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    buffer = io.BytesIO()
    document.save(buffer)
    _atomic_write_bytes(out_path, buffer.getvalue())
