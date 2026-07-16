from unittest.mock import MagicMock

import docx
import httpx
import pytest
from google.genai import errors as genai_errors

import jobs.docx_tailor as docx_tailor_module
from jobs.docx_tailor import (
    MODEL,
    SourceParagraph,
    TailoredResumeEdits,
    _atomic_write_bytes,
    build_tailored_docx,
    estimate_page_risk,
    extract_paragraphs,
    generate_paragraph_edits,
    get_cached_document_stats,
    write_plain_docx,
)


def _make_resume_docx(path):
    document = docx.Document()

    name = document.add_paragraph()
    run = name.add_run("JANE DOE")
    run.bold = True
    run.font.size = docx.shared.Pt(16)

    document.add_paragraph("SUMMARY")

    summary = document.add_paragraph()
    summary.add_run(
        "Software engineer with 5 years building backend systems in Python and Go."
    )

    document.add_paragraph("WORK EXPERIENCE")

    header = document.add_paragraph()
    header_run = header.add_run("Senior Engineer - Acme Corp\tJan 2020 - Present")
    header_run.bold = True

    bullet = document.add_paragraph()
    bullet_run = bullet.add_run("Built and maintained REST APIs serving 1M+ requests/day.")
    bullet_run.font.size = docx.shared.Pt(11)

    document.add_paragraph("GitHub: https://github.com/janedoe")

    document.save(str(path))
    return document


def test_extract_paragraphs_skips_empty_and_keeps_index(tmp_path):
    docx_path = tmp_path / "resume.docx"
    _make_resume_docx(docx_path)

    paragraphs = extract_paragraphs(docx_path)

    assert all(isinstance(p, SourceParagraph) for p in paragraphs)
    assert any(p.text == "JANE DOE" for p in paragraphs)
    assert any("Built and maintained REST APIs" in p.text for p in paragraphs)
    # No blank paragraphs slipped through.
    assert all(p.text.strip() for p in paragraphs)


def test_get_cached_document_stats_parses_real_values(tmp_path):
    docx_path = tmp_path / "resume.docx"
    _make_resume_docx(docx_path)
    # python-docx's blank template ships its own docProps/app.xml (Pages: 1,
    # zeroed word/char counts until Word re-saves it) - confirms the parser
    # reads the real file rather than a hardcoded shape.
    stats = get_cached_document_stats(docx_path)
    assert stats is not None
    assert stats["pages"] == 1


def test_get_cached_document_stats_returns_none_when_file_has_no_app_xml(tmp_path):
    # A .docx-like zip missing docProps/app.xml entirely (e.g. produced by a
    # non-Word tool) must fail gracefully, not raise.
    import zipfile

    bogus_path = tmp_path / "not-quite-a-docx.docx"
    with zipfile.ZipFile(bogus_path, "w") as archive:
        archive.writestr("word/document.xml", "<document/>")

    assert get_cached_document_stats(bogus_path) is None


def test_generate_paragraph_edits_filters_to_rewrite_only():
    expected = TailoredResumeEdits(
        edits=[
            {"index": 0, "action": "keep", "text": None},
            {"index": 1, "action": "rewrite", "text": "Backend engineer specializing in high-throughput Python services."},
        ]
    )
    fake_response = MagicMock(output_text=expected.model_dump_json())
    fake_client = MagicMock()
    fake_client.interactions.create.return_value = fake_response

    paragraphs = [SourceParagraph(index=0, text="JANE DOE"), SourceParagraph(index=1, text="Software engineer...")]
    result = generate_paragraph_edits(paragraphs, "job posting text", "Acme Corp", client=fake_client)

    assert result == {1: "Backend engineer specializing in high-throughput Python services."}
    _, kwargs = fake_client.interactions.create.call_args
    assert kwargs["model"] == MODEL
    assert "job posting text" in kwargs["input"]
    assert "[0] (8 chars) JANE DOE" in kwargs["input"]


def test_generate_paragraph_edits_raises_system_exit_on_api_error():
    fake_client = MagicMock()
    original = genai_errors.APIError(500, {"message": "Internal error", "status": "INTERNAL"})
    fake_client.interactions.create.side_effect = original
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed") as exc_info:
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)
    assert "Internal error" in str(exc_info.value)
    assert exc_info.value.__cause__ is original


def test_generate_paragraph_edits_raises_system_exit_on_network_error():
    fake_client = MagicMock()
    original = httpx.ConnectError("Connection refused")
    fake_client.interactions.create.side_effect = original
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed") as exc_info:
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)
    assert "Connection refused" in str(exc_info.value)
    assert exc_info.value.__cause__ is original


def test_generate_paragraph_edits_raises_system_exit_on_missing_credentials():
    fake_client = MagicMock()
    original = RuntimeError("Could not resolve API token from the environment")
    fake_client.interactions.create.side_effect = original
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed") as exc_info:
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)
    assert "Could not resolve API token" in str(exc_info.value)
    assert exc_info.value.__cause__ is original


def test_generate_paragraph_edits_raises_system_exit_on_unknown_api_response():
    fake_client = MagicMock()
    original = genai_errors.UnknownApiResponseError("Failed to parse response as JSON.")
    fake_client.interactions.create.side_effect = original
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed") as exc_info:
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)
    assert "Failed to parse response as JSON" in str(exc_info.value)
    assert exc_info.value.__cause__ is original


def test_generate_paragraph_edits_raises_system_exit_on_malformed_response():
    fake_response = MagicMock(output_text="{}")
    fake_client = MagicMock()
    fake_client.interactions.create.return_value = fake_response
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed") as exc_info:
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)
    assert "edits" in str(exc_info.value)


def test_generate_paragraph_edits_raises_system_exit_on_non_json_response():
    fake_response = MagicMock(output_text="not valid json at all")
    fake_client = MagicMock()
    fake_client.interactions.create.return_value = fake_response
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed"):
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)


def test_generate_paragraph_edits_raises_system_exit_with_type_name_when_error_message_is_empty():
    fake_client = MagicMock()
    fake_client.interactions.create.side_effect = httpx.HTTPError("")
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(SystemExit, match="Resume paragraph tailoring failed: HTTPError"):
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)


def test_generate_paragraph_edits_does_not_catch_unrelated_exceptions():
    fake_client = MagicMock()
    fake_client.interactions.create.side_effect = KeyError("unexpected")
    paragraphs = [SourceParagraph(index=0, text="JANE DOE")]

    with pytest.raises(KeyError):
        generate_paragraph_edits(paragraphs, "job text", "Acme Corp", client=fake_client)


def test_build_tailored_docx_preserves_formatting_and_changes_text(tmp_path):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)

    paragraphs = extract_paragraphs(source)
    summary_index = next(p.index for p in paragraphs if "Software engineer" in p.text)
    bullet_index = next(p.index for p in paragraphs if "REST APIs" in p.text)

    rewritten = {
        summary_index: "Backend engineer specializing in distributed systems and Python.",
        bullet_index: "Designed REST APIs handling 1M+ requests/day with 99.9% uptime.",
    }

    out_path = tmp_path / "out" / "tailored.docx"
    build_tailored_docx(source, rewritten, out_path)

    assert out_path.exists()
    result_doc = docx.Document(str(out_path))
    texts = [p.text for p in result_doc.paragraphs]
    assert "Backend engineer specializing in distributed systems and Python." in texts
    assert "Designed REST APIs handling 1M+ requests/day with 99.9% uptime." in texts

    # Formatting untouched: name still 16pt bold, bullet's run still 11pt.
    name_paragraph = result_doc.paragraphs[0]
    assert name_paragraph.runs[0].bold is True
    assert name_paragraph.runs[0].font.size == docx.shared.Pt(16)

    bullet_paragraph = next(p for p in result_doc.paragraphs if "Designed REST APIs" in p.text)
    assert bullet_paragraph.runs[0].font.size == docx.shared.Pt(11)

    # The original file on disk must be untouched.
    original_texts = [p.text for p in docx.Document(str(source)).paragraphs]
    assert "Software engineer with 5 years building backend systems in Python and Go." in original_texts


def test_build_tailored_docx_does_not_touch_unlisted_paragraphs(tmp_path):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)
    out_path = tmp_path / "tailored.docx"

    build_tailored_docx(source, {}, out_path)

    result_texts = [p.text for p in docx.Document(str(out_path)).paragraphs]
    original_texts = [p.text for p in docx.Document(str(source)).paragraphs]
    assert result_texts == original_texts


def test_estimate_page_risk_returns_none_without_cached_stats(tmp_path, monkeypatch):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)
    assert estimate_page_risk(source, {}) is None


def test_estimate_page_risk_flags_significant_growth(tmp_path, monkeypatch):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)

    monkeypatch.setattr(
        "jobs.docx_tailor.get_cached_document_stats",
        lambda path: {"pages": 1, "characterswithspaces": 200},
    )

    paragraphs = extract_paragraphs(source)
    bullet_index = next(p.index for p in paragraphs if "REST APIs" in p.text)
    # Replace a ~55-char bullet with a much longer one to force overflow.
    rewritten = {bullet_index: "x" * 400}

    warning = estimate_page_risk(source, rewritten)
    assert warning is not None
    assert "may push past" in warning


def test_estimate_page_risk_silent_when_within_budget(tmp_path, monkeypatch):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)

    monkeypatch.setattr(
        "jobs.docx_tailor.get_cached_document_stats",
        lambda path: {"pages": 2, "characterswithspaces": 5000},
    )

    paragraphs = extract_paragraphs(source)
    bullet_index = next(p.index for p in paragraphs if "REST APIs" in p.text)
    rewritten = {bullet_index: "Shorter bullet."}

    assert estimate_page_risk(source, rewritten) is None


def test_write_plain_docx_creates_one_paragraph_per_block(tmp_path):
    out_path = tmp_path / "cover_letter.docx"
    text = "Dear Hiring Manager,\n\nI'm excited to apply.\n\nBest,\nJane"

    write_plain_docx(text, out_path)

    assert out_path.exists()
    paragraphs = [p.text for p in docx.Document(str(out_path)).paragraphs]
    assert "Dear Hiring Manager," in paragraphs
    assert "I'm excited to apply." in paragraphs


def test_atomic_write_bytes_writes_full_content_and_leaves_no_tmp_file(tmp_path):
    target = tmp_path / "file.docx"

    _atomic_write_bytes(target, b"fake docx bytes")

    assert target.read_bytes() == b"fake docx bytes"
    assert list(tmp_path.glob("*.tmp")) == []


def test_atomic_write_bytes_leaves_prior_content_intact_and_cleans_up_tmp_when_replace_fails(tmp_path, monkeypatch):
    target = tmp_path / "file.docx"
    target.write_bytes(b"original docx bytes")

    monkeypatch.setattr(docx_tailor_module.os, "replace", MagicMock(side_effect=OSError("simulated disk failure")))

    with pytest.raises(OSError, match="simulated disk failure"):
        _atomic_write_bytes(target, b"new bytes that must never land")

    assert target.read_bytes() == b"original docx bytes"
    assert list(tmp_path.glob("*.tmp")) == []


def test_build_tailored_docx_leaves_no_partial_file_when_replace_fails(tmp_path, monkeypatch):
    source = tmp_path / "resume.docx"
    _make_resume_docx(source)
    out_path = tmp_path / "out" / "tailored.docx"

    monkeypatch.setattr(docx_tailor_module.os, "replace", MagicMock(side_effect=OSError("simulated disk failure")))

    with pytest.raises(OSError, match="simulated disk failure"):
        build_tailored_docx(source, {}, out_path)

    assert not out_path.exists()
    assert list((tmp_path / "out").glob("*.tmp")) == []
