import io
from unittest.mock import MagicMock

import docx
import httpx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google.genai._gaos.lib import compat_errors

from resume.extract import (
    MODEL,
    ResumeProfile,
    _iter_body_content_including_sdt,
    _iter_row_tc_elements,
    _table_has_any_sdt_wrapping,
    extract_profile,
    extract_text_from_docx,
)


def _wrap_in_sdt(element):
    """Wrap an existing paragraph/table oxml element in a
    `<w:sdt><w:sdtContent>...</w:sdtContent></w:sdt>` content-control
    wrapper, replacing it in place at the same document position -
    simulates how Word's own resume/CV templates mark up a section as a
    content control. Returns the new `w:sdt` element, so callers can wrap
    it again to build a nested chain."""
    parent = element.getparent()
    index = list(parent).index(element)
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)
    parent.remove(element)
    sdt_content.append(element)
    parent.insert(index, sdt)
    return sdt


def test_extract_profile_calls_gemini_with_expected_shape_and_parses_output():
    expected = ResumeProfile(
        full_name="Jane Doe",
        years_experience=5.0,
        seniority="senior",
        core_skills=["Python", "PyTorch", "LangGraph"],
        domains=["Generative AI", "NLP"],
        past_roles=["Senior ML Engineer at Acme"],
        summary="Senior ML engineer with 5 years building production LLM systems.",
    )
    fake_response = MagicMock(output_text=expected.model_dump_json())
    fake_client = MagicMock()
    fake_client.interactions.create.return_value = fake_response

    result = extract_profile("some raw resume text", client=fake_client)

    assert result == expected
    _, kwargs = fake_client.interactions.create.call_args
    assert kwargs["model"] == MODEL
    assert kwargs["input"] == "some raw resume text"
    assert kwargs["response_format"]["schema"] == ResumeProfile.model_json_schema()


def test_extract_profile_raises_system_exit_on_connection_error():
    fake_client = MagicMock()
    original = compat_errors.APIConnectionError(
        message="Connection refused", request=httpx.Request("POST", "https://example.com")
    )
    fake_client.interactions.create.side_effect = original

    with pytest.raises(SystemExit, match="Resume profile extraction failed: Connection refused") as exc_info:
        extract_profile("some raw resume text", client=fake_client)
    assert exc_info.value.__cause__ is original


def test_extract_text_from_docx_joins_paragraph_text_with_newlines():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior ML Engineer with 5 years of experience.")
    doc.add_paragraph("Skills: Python, PyTorch, LangGraph")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == (
        "Jane Doe\n"
        "Senior ML Engineer with 5 years of experience.\n"
        "Skills: Python, PyTorch, LangGraph"
    )


def test_extract_text_from_docx_includes_table_content_in_document_order():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Years"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "5"
    doc.add_paragraph("References available on request.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == (
        "Jane Doe\n"
        "Skill | Years\n"
        "Python | 5\n"
        "References available on request."
    )


def test_extract_text_from_docx_skips_all_blank_table_rows():
    # An unfilled template table (grid laid out, nothing typed in) must not
    # inject a stray " | " separator into the output - the caller in
    # views/admin.py treats a blank-after-strip result as "no text found".
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_table(rows=2, cols=2)  # every cell left empty

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Jane Doe"


def test_extract_text_from_docx_dedupes_merged_cell_text():
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Skills"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "5 years"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Skills\nPython | 5 years"


def test_extract_text_from_docx_flattens_newlines_inside_a_table_cell():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    cell.text = "Python"
    cell.add_paragraph("PyTorch")
    table.cell(0, 1).text = "5 years"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Python PyTorch | 5 years"


def test_extract_text_from_docx_reads_a_paragraph_wrapped_in_a_content_control():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    wrapped = doc.add_paragraph("Senior ML Engineer")
    doc.add_paragraph("References available on request.")
    _wrap_in_sdt(wrapped._p)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Jane Doe\nSenior ML Engineer\nReferences available on request."


def test_extract_text_from_docx_reads_a_table_wrapped_in_a_content_control():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "5 years"
    doc.add_paragraph("References available on request.")
    _wrap_in_sdt(table._tbl)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Jane Doe\nPython | 5 years\nReferences available on request."


def test_extract_text_from_docx_reads_content_nested_two_levels_deep_in_content_controls():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    wrapped = doc.add_paragraph("Senior ML Engineer")
    inner_sdt = _wrap_in_sdt(wrapped._p)
    _wrap_in_sdt(inner_sdt)  # a content control nested inside another one

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Jane Doe\nSenior ML Engineer"


def test_extract_text_from_docx_skips_content_past_the_depth_cap_without_crashing():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    buried = doc.add_paragraph("Buried past the depth cap")
    doc.add_paragraph("References available on request.")

    element = buried._p
    for _ in range(30):  # well past _MAX_SDT_DEPTH (20)
        element = _wrap_in_sdt(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)  # must not raise (no RecursionError, no hang)

    assert "Buried past the depth cap" not in result
    assert result == "Jane Doe\nReferences available on request."


def test_iter_body_content_including_sdt_does_not_recursion_error_on_a_pathologically_deep_chain():
    # Exercises the walk directly, in-memory (no save/reload round-trip):
    # lxml's XML *parser* itself refuses to load a document nested past
    # ~256 levels ("Excessive depth in document"), which would mask
    # whether our own walk is actually recursion-safe - building the tree
    # via the oxml element API instead (as `_wrap_in_sdt` does) isn't
    # subject to that parse-time guard, so this depth is only reachable
    # from a document already loaded via some other means (e.g. a parser
    # configured with XML_PARSE_HUGE), which is exactly the untrusted-input
    # case `_MAX_SDT_DEPTH` defends against.
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    buried = doc.add_paragraph("Buried impossibly deep")

    element = buried._p
    for _ in range(1500):  # comfortably past Python's default recursion limit (1000)
        element = _wrap_in_sdt(element)

    items = list(_iter_body_content_including_sdt(doc))  # must not raise RecursionError

    assert [item.text for item in items] == ["Jane Doe"]


def test_extract_text_from_docx_reads_content_at_exactly_the_depth_cap():
    doc = docx.Document()
    p = doc.add_paragraph("At the cap")

    element = p._p
    for _ in range(20):  # exactly _MAX_SDT_DEPTH - must still be read
        element = _wrap_in_sdt(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "At the cap"


def test_extract_text_from_docx_skips_content_one_level_past_the_depth_cap():
    doc = docx.Document()
    p = doc.add_paragraph("Just past the cap")

    element = p._p
    for _ in range(21):  # one more than _MAX_SDT_DEPTH - must be dropped
        element = _wrap_in_sdt(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == ""


def test_iter_body_content_including_sdt_skips_an_sdt_with_no_sdt_content_child():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    sdt = OxmlElement("w:sdt")  # no <w:sdtContent> child at all - malformed but must not crash
    doc.element.body.insert(0, sdt)
    doc.add_paragraph("References available on request.")

    items = list(_iter_body_content_including_sdt(doc))

    assert [item.text for item in items] == ["Jane Doe", "References available on request."]


def test_iter_body_content_including_sdt_handles_an_empty_sdt_content():
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    sdt = OxmlElement("w:sdt")
    sdt.append(OxmlElement("w:sdtContent"))  # present but zero children
    doc.element.body.insert(0, sdt)

    items = list(_iter_body_content_including_sdt(doc))

    assert [item.text for item in items] == ["Jane Doe"]


def test_extract_text_from_docx_preserves_order_across_multiple_interleaved_wraps():
    doc = docx.Document()
    doc.add_paragraph("Plain 1")
    wrapped_p = doc.add_paragraph("Wrapped paragraph")
    _wrap_in_sdt(wrapped_p._p)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Plain table"
    doc.add_paragraph("Plain 2")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Plain 1\nWrapped paragraph\nPlain table\nPlain 2"


def test_iter_body_content_including_sdt_reads_multiple_block_children_in_one_wrapper():
    doc = docx.Document()
    first = doc.add_paragraph("First in wrapper")
    second = doc.add_paragraph("Second in wrapper")

    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)
    body = doc.element.body
    index = list(body).index(first._p)
    body.remove(first._p)
    body.remove(second._p)
    sdt_content.append(first._p)
    sdt_content.append(second._p)
    body.insert(index, sdt)

    items = list(_iter_body_content_including_sdt(doc))

    assert [item.text for item in items] == ["First in wrapper", "Second in wrapper"]


def test_extract_text_from_docx_reads_a_table_row_wrapped_in_a_content_control():
    doc = docx.Document()
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "Row 1"
    table.cell(1, 0).text = "Row 2 (wrapped)"
    table.cell(2, 0).text = "Row 3"
    _wrap_in_sdt(table.rows[1]._tr)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Row 1\nRow 2 (wrapped)\nRow 3"


def test_extract_text_from_docx_reads_a_cell_wrapped_in_a_content_control():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B (wrapped)"
    table.cell(0, 2).text = "C"
    _wrap_in_sdt(table.cell(0, 1)._tc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "A | B (wrapped) | C"


def test_extract_text_from_docx_reads_an_in_cell_paragraph_wrapped_in_a_content_control():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "First paragraph"
    wrapped_paragraph = cell.add_paragraph("Second paragraph (wrapped)")
    _wrap_in_sdt(wrapped_paragraph._p)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "First paragraph Second paragraph (wrapped)"


def test_extract_text_from_docx_still_dedupes_a_horizontal_merge_when_the_row_is_wrapped():
    # Combines CAP-1 with the existing merge-dedup guarantee (CAP-4) - a
    # merged cell must still appear once, not twice, even when the whole
    # row containing it is itself sdt-wrapped.
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Skills (wrapped row)"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "5 years"
    _wrap_in_sdt(table.rows[0]._tr)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Skills (wrapped row)\nPython | 5 years"


def test_extract_text_from_docx_still_repeats_a_vertical_merge_when_nothing_nearby_is_wrapped():
    # Pins CAP-4's existing vertical-merge-repetition guarantee completely
    # untouched by this diff (no sdt-wrapping anywhere in or near the
    # table) - the counterpart to the two residual tests below, which cover
    # what happens once sdt-wrapping IS present nearby.
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Skills | A\nSkills | B"


def test_iter_row_tc_elements_shows_blank_instead_of_wrong_content_when_the_top_cell_is_wrapped():
    # A more serious trap than a wrapped continuation cell (see the test
    # below): wrapping the TOP (content-holding) cell of a vertical merge
    # doesn't just make it unreachable - CT_Tc.grid_offset's search inside
    # the row above (tc_at_grid_offset, direct-<w:tc>-children only) can
    # silently resolve to the WRONG cell instead of raising, since wrapping
    # removes the real cell from that row's direct-child count. Confirmed
    # via a live repro before this guard existed: the continuation row
    # showed the *sibling* cell's text ("A") instead of "Skills" or a
    # blank. The fix must degrade to blank, never to wrong-but-plausible
    # content - this is why any w:sdt use anywhere in the table disables
    # vMerge resolution for the WHOLE table (see _table_has_any_sdt_wrapping),
    # not just a narrower per-row check (an earlier, rejected design - see
    # the row-level tests below for why that turned out unsafe too).
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"
    _wrap_in_sdt(table.cell(0, 0)._tc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)  # must not raise, and must not show "A" for column 0

    assert result == "Skills | A\n | B"


def test_iter_row_tc_elements_shows_its_own_literal_content_for_a_wrapped_vmerge_continuation_cell():
    # Whitebox test of _iter_row_tc_elements in isolation, given an
    # explicit vmerge_resolution_is_safe=False - does NOT exercise
    # _table_has_any_sdt_wrapping itself (that's covered by the
    # extract_text_from_docx-level tests around this one). What this DOES
    # prove: even setting the table-level flag aside, a cell found by
    # unwrapping a w:sdt still never reaches vMerge resolution, because
    # its real parent is w:sdtContent, not w:tr - rather than resolve
    # incorrectly, it's shown with its own (here, empty) literal content -
    # and critically, this must not corrupt or crash extraction of the
    # row's OTHER cells.
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"
    # table.cell(1, 0) resolves vMerge automatically and returns the SAME
    # _Cell as cell(0, 0) - the raw, physically-present continuation <w:tc>
    # (with its own empty <w:p/>) must be found via the row's own oxml
    # children instead.
    continuation_tc = table.rows[1]._tr.tc_lst[0]
    assert continuation_tc.vMerge == "continue"
    _wrap_in_sdt(continuation_tc)

    tc_elements = list(_iter_row_tc_elements(table.rows[1]._tr, vmerge_resolution_is_safe=False))

    assert len(tc_elements) == 2
    assert tc_elements[0] is continuation_tc  # shown literally, not resolved to the cell above
    assert tc_elements[1].xpath("string(.)") == "B"  # the sibling cell is unaffected

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)  # must not raise

    assert result == "Skills | A\n | B"


def test_extract_text_from_docx_does_not_crash_when_a_vmerge_continuation_row_is_wrapped():
    # A real bug found during review, not anticipated up front: CT_Tc._tr_above's
    # xpath ("./ancestor::w:tr[position()=1]/preceding-sibling::w:tr[1]") raises
    # ValueError ("no tr above topmost tr in w:tbl") when the row it's
    # searching from is itself w:sdt-wrapped, since that row's real parent
    # becomes w:sdtContent, not w:tbl - preceding-sibling then finds nothing.
    # A per-row guard checking only the CELLS (not the row element itself)
    # didn't catch this; confirmed via a live repro that it crashed the
    # entire extraction, not just this one row.
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"
    _wrap_in_sdt(table.rows[1]._tr)  # wrap the row holding the continuation cell

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)  # must not raise

    assert result == "Skills | A\n | B"


def test_extract_text_from_docx_does_not_crash_when_the_row_above_a_vmerge_continuation_is_wrapped():
    # The mirror image of the test above: wrapping the row ABOVE the
    # continuation cell (the one holding the real content) also breaks
    # _tr_above's xpath assumptions for the continuation row's own lookup,
    # for the identical underlying reason.
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"
    _wrap_in_sdt(table.rows[0]._tr)  # wrap the row holding the real content

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)  # must not raise

    assert result == "Skills | A\n | B"


def test_extract_text_from_docx_disables_vmerge_repetition_for_an_unrelated_merge_in_the_same_table():
    # Demonstrates the actual documented tradeoff, not just asserting it in
    # prose: wrapping a cell involved in ONE vertical merge (column 0)
    # disables repetition for a SECOND, structurally independent vertical
    # merge elsewhere in the SAME table (column 1) that has nothing to do
    # with the wrapped cell - the whole-table guard is intentionally this
    # wide, in exchange for the correctness guarantees the earlier
    # per-row/per-cell guards couldn't provide.
    doc = docx.Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(2, 1).merge(table.cell(3, 1))
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "X"
    table.cell(1, 1).text = "Y"
    table.cell(2, 0).text = "P"
    table.cell(2, 1).text = "Notes"
    table.cell(3, 0).text = "Q"
    _wrap_in_sdt(table.cell(0, 0)._tc)  # only touches the FIRST merge's top cell

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text_from_docx(buffer)

    assert result == "Skills | X\n | Y\nP | Notes\nQ | "


def test_table_has_any_sdt_wrapping_is_false_for_a_table_with_no_rows():
    # A <w:tbl> with zero <w:tr> children is invalid per OOXML but
    # plausible from a hand-crafted or corrupted file - must degrade
    # safely (no IndexError/ZeroDivisionError) rather than crash the
    # detector for the whole document.
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    tbl_element = table._tbl
    for tr in list(tbl_element.findall(qn("w:tr"))):
        tbl_element.remove(tr)

    assert _table_has_any_sdt_wrapping(tbl_element) is False
