"""Tests for `jobs.ui_actions` - job-scoped wrappers shared by every
Streamlit page. `draft_and_save_outreach` had zero test coverage before
this file (the write-atomicity/error-message fix in this diff is what
first required it)."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

import docx
import pytest
from streamlit.testing.v1 import AppTest

import jobs.cli as jobs_cli
import jobs.ui_actions as ui_actions
from jobs.cli import DEFAULT_GENERATED_CV_DIR, _tailored_docx_paths
from jobs.db import connect as connect_jobs
from jobs.db import insert_job
from jobs.extract import JobExtraction
from jobs.outreach import LINKEDIN_NOTE, OutreachDraft
from jobs.outreach_db import ensure_schema as ensure_outreach_schema
from jobs.outreach_db import list_outreach_messages
from jobs.tailor import TailoredApplication
from jobs.ui_actions import (
    _outreach_message_path,
    draft_and_save_outreach,
    error_display_text,
    generate_tailored_docx_for_job,
)
from resume.db import connect as connect_profile
from resume.db import insert_narrative, insert_profile
from resume.extract import ResumeProfile

RESUME_TEXT = "Jane Doe - Senior Backend Engineer. 5 years Python."

# Resolved at import time (before any test's chdir into a tmp sandbox) so
# AppTest.from_file can find the real scripts regardless of the test's cwd.
_VIEWS_DIR = Path(__file__).resolve().parent.parent / "views"
INTAKE_PY = str(_VIEWS_DIR / "intake.py")
JOBS_LIST_PY = str(_VIEWS_DIR / "jobs_list.py")


def _seed_profile_and_narrative(profile_db_path: Path) -> None:
    conn = connect_profile(profile_db_path)
    try:
        insert_profile(
            conn,
            RESUME_TEXT,
            ResumeProfile(
                full_name="Jane Doe",
                years_experience=5,
                seniority="Senior",
                core_skills=["Python"],
                domains=["Backend"],
                past_roles=["Engineer"],
                summary="Backend engineer.",
            ),
        )
        insert_narrative(conn, "Why AI, why UK, why them - the candidate's narrative core.")
    finally:
        conn.close()


@pytest.fixture
def ui_outreach_env(tmp_path, monkeypatch):
    jobs_db = tmp_path / "jobs.db"
    profile_db = tmp_path / "profile.db"

    conn = connect_jobs(jobs_db)
    try:
        job_id = insert_job(
            conn,
            "job posting text",
            JobExtraction(
                job_title="AI Engineer",
                company_name="Bending Spoons",
                is_agency_posting=False,
                recruiter_name="Sarah Cole",
                recruiter_contact="sarah@bendingspoons.com",
            ),
        )
    finally:
        conn.close()

    _seed_profile_and_narrative(profile_db)

    monkeypatch.setattr(
        ui_actions, "draft_outreach_message", MagicMock(return_value=OutreachDraft(message="Hi Sarah, let's talk."))
    )

    return {"jobs_db": jobs_db, "profile_db": profile_db, "job_id": job_id}


def test_draft_and_save_outreach_writes_the_expected_txt_file_and_stores_no_message_column(ui_outreach_env):
    draft = draft_and_save_outreach(
        ui_outreach_env["job_id"],
        LINKEDIN_NOTE,
        contact_id=None,
        contact_name="Sarah Cole",
        contact_title="Recruiter",
        purpose=None,
        jobs_db=str(ui_outreach_env["jobs_db"]),
        profile_db=str(ui_outreach_env["profile_db"]),
    )

    assert draft.message == "Hi Sarah, let's talk."

    conn = connect_jobs(ui_outreach_env["jobs_db"])
    try:
        ensure_outreach_schema(conn)
        messages = list_outreach_messages(conn, ui_outreach_env["job_id"])
    finally:
        conn.close()

    assert len(messages) == 1
    message_row = messages[0]
    assert "message" not in message_row.keys()

    path = _outreach_message_path(
        "Bending Spoons", ui_outreach_env["job_id"], LINKEDIN_NOTE, message_row["id"], ui_actions.DEFAULT_GENERATED_CV_DIR
    )
    assert path.exists()
    assert path.read_text(encoding="utf-8") == "Hi Sarah, let's talk."


def test_draft_and_save_outreach_write_failure_after_db_commit_raises_error_containing_the_drafted_text(
    ui_outreach_env, monkeypatch
):
    """The DB row commits before the file write is attempted, so a write
    failure here is worse than an ordinary I/O error - unlike the CLI's
    equivalent fix, the recovery text must be *inside* the exception
    message, since a Streamlit user never sees server-side print() output."""
    monkeypatch.setattr(ui_actions, "_atomic_write_text", MagicMock(side_effect=OSError("disk full")))

    with pytest.raises(SystemExit) as exc_info:
        draft_and_save_outreach(
            ui_outreach_env["job_id"],
            LINKEDIN_NOTE,
            contact_id=None,
            contact_name="Sarah Cole",
            contact_title="Recruiter",
            purpose=None,
            jobs_db=str(ui_outreach_env["jobs_db"]),
            profile_db=str(ui_outreach_env["profile_db"]),
        )

    message = str(exc_info.value)
    assert "was logged to the database" in message
    assert "Hi Sarah, let's talk." in message

    # The DB row is still there (insert-only convention, no rollback) - the
    # exact orphaned-metadata state the error message must warn about.
    conn = connect_jobs(ui_outreach_env["jobs_db"])
    try:
        ensure_outreach_schema(conn)
        messages = list_outreach_messages(conn, ui_outreach_env["job_id"])
    finally:
        conn.close()
    assert len(messages) == 1


def test_error_display_text_returns_str_exc_unchanged_when_non_empty():
    exc = ValueError("disk full")
    assert error_display_text(exc) == "disk full"


def test_error_display_text_falls_back_to_type_name_when_str_exc_is_empty():
    exc = OSError()
    assert str(exc) == ""

    text = error_display_text(exc)

    assert text != ""
    assert "OSError" in text


def test_error_display_text_falls_back_to_type_name_when_str_exc_is_whitespace_only():
    exc = ValueError("   ")

    text = error_display_text(exc)

    assert text.strip() != ""
    assert "ValueError" in text


def test_error_display_text_handles_system_exit_unchanged_when_non_empty():
    # SystemExit is what most real call sites (views/intake.py, views/jobs_list.py)
    # actually pass - it's a BaseException, not an Exception subclass.
    exc = SystemExit("No job #99 found in data/jobs.db")
    assert error_display_text(exc) == "No job #99 found in data/jobs.db"


def test_error_display_text_survives_a_str_that_itself_raises():
    class BrokenStr(Exception):
        def __str__(self):
            raise RuntimeError("__str__ is broken")

    text = error_display_text(BrokenStr())

    assert text != ""
    assert "BrokenStr" in text


def _make_source_resume_docx(path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("JANE DOE")
    document.add_paragraph("SUMMARY")
    document.add_paragraph("Software engineer with 5 years building backend systems in Python.")
    document.save(str(path))


@pytest.fixture
def ui_tailor_env(tmp_path, monkeypatch):
    """A ready-to-use jobs.db (one job) + profile.db + a real source .docx,
    with the two LLM-calling functions the docx path uses monkeypatched -
    same mocking approach as test_jobs_cli.py's `env` fixture (independent
    setup, not shared, to keep this file's fixtures self-contained), scoped
    to generate_tailored_docx_for_job instead of the CLI wiring."""
    jobs_db = tmp_path / "jobs.db"
    profile_db = tmp_path / "profile.db"
    resume_dir = tmp_path / "my-resume"
    resume_dir.mkdir()
    _make_source_resume_docx(resume_dir / "resume.docx")

    conn = connect_jobs(jobs_db)
    try:
        job_id = insert_job(
            conn,
            "job posting text",
            JobExtraction(job_title="AI Engineer", company_name="Acme AI", is_agency_posting=False),
        )
    finally:
        conn.close()

    _seed_profile_and_narrative(profile_db)

    tailor_calls = MagicMock(
        return_value=TailoredApplication(
            tailored_resume="TAILORED RESUME",
            cover_letter="COVER LETTER",
            evidence_notes=["evidence note"],
            portfolio_gaps=["portfolio gap"],
        )
    )
    paragraph_calls = MagicMock(return_value={})
    monkeypatch.setattr(jobs_cli, "generate_tailored_application", tailor_calls)
    monkeypatch.setattr(jobs_cli, "generate_paragraph_edits", paragraph_calls)
    monkeypatch.setattr(ui_actions, "DEFAULT_SOURCE_RESUME_DIR", str(resume_dir))
    monkeypatch.setattr(ui_actions, "DEFAULT_GENERATED_CV_DIR", str(tmp_path / "generated_cv"))

    return {
        "jobs_db": jobs_db,
        "profile_db": profile_db,
        "job_id": job_id,
        "tailor_calls": tailor_calls,
        "paragraph_calls": paragraph_calls,
    }


def test_generate_tailored_docx_for_job_with_force_true_bypasses_cache_and_calls_llm_again(ui_tailor_env):
    out_dir, _ = generate_tailored_docx_for_job(
        ui_tailor_env["job_id"], str(ui_tailor_env["jobs_db"]), str(ui_tailor_env["profile_db"])
    )
    cover_letter_path = out_dir / f"{ui_tailor_env['job_id']}_cover_letter.docx"
    first_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text
    assert ui_tailor_env["tailor_calls"].call_count == 1

    ui_tailor_env["tailor_calls"].return_value = TailoredApplication(
        tailored_resume="TAILORED RESUME_V2",
        cover_letter="COVER LETTER_V2",
        evidence_notes=["evidence note_V2"],
        portfolio_gaps=["portfolio gap_V2"],
    )
    generate_tailored_docx_for_job(
        ui_tailor_env["job_id"], str(ui_tailor_env["jobs_db"]), str(ui_tailor_env["profile_db"]), force=True
    )

    assert ui_tailor_env["tailor_calls"].call_count == 2
    assert ui_tailor_env["paragraph_calls"].call_count == 2
    second_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text
    assert second_cover_text != first_cover_text
    assert "COVER LETTER_V2" in second_cover_text


def test_generate_tailored_docx_for_job_with_force_false_default_still_uses_cache(ui_tailor_env):
    out_dir, _ = generate_tailored_docx_for_job(
        ui_tailor_env["job_id"], str(ui_tailor_env["jobs_db"]), str(ui_tailor_env["profile_db"])
    )
    cover_letter_path = out_dir / f"{ui_tailor_env['job_id']}_cover_letter.docx"
    first_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text
    assert ui_tailor_env["tailor_calls"].call_count == 1

    ui_tailor_env["tailor_calls"].return_value = TailoredApplication(
        tailored_resume="TAILORED RESUME_V2",
        cover_letter="COVER LETTER_V2",
        evidence_notes=["evidence note_V2"],
        portfolio_gaps=["portfolio gap_V2"],
    )
    generate_tailored_docx_for_job(ui_tailor_env["job_id"], str(ui_tailor_env["jobs_db"]), str(ui_tailor_env["profile_db"]))

    assert ui_tailor_env["tailor_calls"].call_count == 1
    assert ui_tailor_env["paragraph_calls"].call_count == 1
    second_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text
    assert second_cover_text == first_cover_text


def _seed_job_with_existing_docx_cache(jobs_db_path) -> int:
    """Insert a job AND touch its docx cache files (only existence is
    checked by `already_generated`, not content) so the tailor button
    renders as "Regenerate", not "Generate"."""
    conn = connect_jobs(jobs_db_path)
    try:
        job_id = insert_job(
            conn,
            "job posting text",
            JobExtraction(job_title="AI Engineer", company_name="Bending Spoons", is_agency_posting=False),
        )
    finally:
        conn.close()

    resume_path, cover_letter_path = _tailored_docx_paths("Bending Spoons", job_id, DEFAULT_GENERATED_CV_DIR)
    resume_path.parent.mkdir(parents=True, exist_ok=True)
    resume_path.touch()
    cover_letter_path.touch()
    return job_id


def test_jobs_list_regenerate_button_passes_force_true_when_docx_cache_already_exists(streamlit_data_env, monkeypatch):
    """View-wiring check: the force=True/False tests above only call
    generate_tailored_docx_for_job directly - they can't catch a regression
    at the actual views/jobs_list.py call site (e.g. reverting to a
    hardcoded force=False, or inverting the boolean). This drives the real
    page via AppTest instead."""
    job_id = _seed_job_with_existing_docx_cache(streamlit_data_env["jobs_db"])
    mock_tailor = MagicMock(side_effect=SystemExit())
    monkeypatch.setattr("jobs.ui_actions.generate_tailored_docx_for_job", mock_tailor)

    at = AppTest.from_file(JOBS_LIST_PY)
    at.run()

    tailor_button = next(b for b in at.button if b.label == "Regenerate tailored resume & cover letter")
    tailor_button.click().run()

    assert mock_tailor.call_args.kwargs["force"] is True


def test_intake_regenerate_button_passes_force_true_when_docx_cache_already_exists(streamlit_data_env, monkeypatch):
    job_id = _seed_job_with_existing_docx_cache(streamlit_data_env["jobs_db"])
    mock_tailor = MagicMock(side_effect=SystemExit())
    monkeypatch.setattr("jobs.ui_actions.generate_tailored_docx_for_job", mock_tailor)

    at = AppTest.from_file(INTAKE_PY)
    at.session_state["extraction"] = JobExtraction(
        job_title="AI Engineer", company_name="Bending Spoons", is_agency_posting=False
    )
    at.session_state["resolved_employer"] = "Bending Spoons"
    at.session_state["saved_job_id"] = job_id
    at.run()

    tailor_button = next(b for b in at.button if "tailored resume" in b.label)
    tailor_button.click().run()

    assert mock_tailor.call_args.kwargs["force"] is True
