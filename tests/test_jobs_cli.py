"""Tests for `jobs.cli`'s tailoring + legacy-migration commands.

Exercised through `build_parser()`/`args.func(args)` (the real CLI wiring),
not only the internal helper functions - the previous implementation
attempt's test gap (testing helpers directly) is exactly what let a
caching bug in `_cmd_tailor`'s wiring through undetected. See
`_bmad-output/implementation-artifacts/spec-tailored-content-file-only-storage.md`
Spec Change Log for the full story.
"""

from __future__ import annotations

import builtins
import sqlite3
from pathlib import Path
from unittest.mock import MagicMock

import docx
import pytest

import jobs.cli as jobs_cli
from jobs.cli import _outreach_message_path, _sanitize_filename, _tailored_docx_paths, build_parser
from jobs.db import connect as connect_jobs
from jobs.db import get_job, insert_job, try_claim_tailoring_lock
from jobs.extract import JobExtraction
from jobs.outreach import EMAIL, LINKEDIN_NOTE, OutreachDraft
from jobs.outreach_db import ensure_schema as ensure_outreach_schema
from jobs.outreach_db import list_outreach_messages
from jobs.tailor import TailoredApplication
from resume.db import connect as connect_profile
from resume.db import insert_narrative, insert_profile
from resume.extract import ResumeProfile

RESUME_TEXT = "Jane Doe - Senior Backend Engineer. 5 years Python. No GitHub link in this resume."


def _add_profile_version(profile_db_path: Path, resume_text: str) -> None:
    conn = connect_profile(profile_db_path)
    try:
        insert_profile(
            conn,
            resume_text,
            ResumeProfile(
                full_name="Jane Doe",
                years_experience=5,
                seniority="Senior",
                core_skills=["Python"],
                domains=["Backend"],
                past_roles=["Engineer"],
                summary="Backend engineer.",
            ),
        )
    finally:
        conn.close()


def _make_job(jobs_db_path: Path, *, company_name="Acme AI", raw_text="job posting text") -> int:
    conn = connect_jobs(jobs_db_path)
    try:
        job_id = insert_job(
            conn, raw_text, JobExtraction(job_title="AI Engineer", company_name=company_name, is_agency_posting=False)
        )
    finally:
        conn.close()
    return job_id


def _make_source_resume_docx(path: Path) -> None:
    document = docx.Document()
    document.add_paragraph("JANE DOE")
    document.add_paragraph("SUMMARY")
    document.add_paragraph("Software engineer with 5 years building backend systems in Python.")
    document.save(str(path))


def _fake_tailored_application(suffix: str = "") -> TailoredApplication:
    return TailoredApplication(
        tailored_resume=f"TAILORED RESUME{suffix}",
        cover_letter=f"COVER LETTER{suffix}",
        evidence_notes=[f"evidence note{suffix}"],
        portfolio_gaps=[f"portfolio gap{suffix}"],
    )


def _run(argv: list[str]) -> None:
    args = build_parser().parse_args(argv)
    args.func(args)


def _add_legacy_columns_with_text(jobs_db_path: Path, job_id: int, tailored_resume: str, cover_letter: str) -> None:
    """Simulate a jobs.db created before this refactor: `tailored_resume`/
    `cover_letter` still exist as real columns with real data, since
    `_ensure_columns()` only ever adds columns, never drops them. A fresh
    `connect()` no longer creates these columns at all (they're not in
    SCHEMA anymore), so this reaches for raw sqlite3 to simulate the
    pre-existing state directly."""
    conn = sqlite3.connect(jobs_db_path)
    conn.execute("ALTER TABLE jobs ADD COLUMN tailored_resume TEXT")
    conn.execute("ALTER TABLE jobs ADD COLUMN cover_letter TEXT")
    conn.execute(
        "UPDATE jobs SET tailored_resume = ?, cover_letter = ? WHERE id = ?", (tailored_resume, cover_letter, job_id)
    )
    conn.commit()
    conn.close()


@pytest.fixture
def env(tmp_path, monkeypatch):
    """A ready-to-use jobs.db (one job) + profile.db + a real source .docx,
    with the two LLM-calling functions the docx path uses monkeypatched so
    tests can assert on call counts without hitting a real API. No GitHub
    URL appears in the resume text, so `_fetch_github_evidence` also never
    makes a real network call."""
    jobs_db = tmp_path / "jobs.db"
    profile_db = tmp_path / "profile.db"
    resume_dir = tmp_path / "my-resume"
    out_dir = tmp_path / "generated_cv"
    resume_dir.mkdir()

    job_id = _make_job(jobs_db)
    _add_profile_version(profile_db, RESUME_TEXT)
    _make_source_resume_docx(resume_dir / "resume.docx")

    tailor_calls = MagicMock(return_value=_fake_tailored_application())
    paragraph_calls = MagicMock(return_value={})
    monkeypatch.setattr(jobs_cli, "generate_tailored_application", tailor_calls)
    monkeypatch.setattr(jobs_cli, "generate_paragraph_edits", paragraph_calls)

    return {
        "jobs_db": jobs_db,
        "profile_db": profile_db,
        "resume_dir": resume_dir,
        "out_dir": out_dir,
        "job_id": job_id,
        "tailor_calls": tailor_calls,
        "paragraph_calls": paragraph_calls,
    }


def _tailor_docx_argv(env, job_id=None, *, force: bool = False) -> list[str]:
    argv = [
        "tailor-docx",
        str(job_id if job_id is not None else env["job_id"]),
        "--db",
        str(env["jobs_db"]),
        "--profile-db",
        str(env["profile_db"]),
        "--resume-dir",
        str(env["resume_dir"]),
        "--out-dir",
        str(env["out_dir"]),
    ]
    if force:
        argv.append("--force")
    return argv


def _get_job_row(jobs_db_path, job_id):
    conn = connect_jobs(jobs_db_path)
    try:
        return get_job(conn, job_id)
    finally:
        conn.close()


# --------------------------------------------------------------------------
# `tailor-docx` - job_id-keyed docx cache-check
# --------------------------------------------------------------------------


def test_tailor_docx_first_run_generates_via_llm_and_writes_job_id_keyed_files(env):
    _run(_tailor_docx_argv(env))

    resume_path, cover_letter_path = _tailored_docx_paths("Acme AI", env["job_id"], str(env["out_dir"]))
    assert resume_path.exists()
    assert cover_letter_path.exists()
    assert env["tailor_calls"].call_count == 1
    assert env["paragraph_calls"].call_count == 1

    row = _get_job_row(env["jobs_db"], env["job_id"])
    assert row["tailor_hash"] is not None
    assert row["tailor_evidence_notes"] == '["evidence note"]'
    assert row["tailor_portfolio_gaps"] == '["portfolio gap"]'


def test_tailor_docx_rerun_with_nothing_changed_is_a_cache_hit_with_no_llm_call(env):
    _run(_tailor_docx_argv(env))
    resume_path, _ = _tailored_docx_paths("Acme AI", env["job_id"], str(env["out_dir"]))
    mtime_before = resume_path.stat().st_mtime

    _run(_tailor_docx_argv(env))

    assert env["tailor_calls"].call_count == 1
    assert env["paragraph_calls"].call_count == 1
    assert resume_path.stat().st_mtime == mtime_before


def test_tailor_docx_second_job_same_company_gets_its_own_files_and_does_not_touch_the_first(env):
    _run(_tailor_docx_argv(env))
    first_resume_path, _ = _tailored_docx_paths("Acme AI", env["job_id"], str(env["out_dir"]))
    first_content = first_resume_path.read_bytes()

    second_job_id = _make_job(env["jobs_db"], company_name="Acme AI", raw_text="a different job posting entirely")
    _run(_tailor_docx_argv(env, job_id=second_job_id))

    second_resume_path, second_cover_path = _tailored_docx_paths("Acme AI", second_job_id, str(env["out_dir"]))
    assert second_resume_path.exists()
    assert second_cover_path.exists()
    assert second_resume_path != first_resume_path
    assert first_resume_path.exists()
    assert first_resume_path.read_bytes() == first_content
    assert env["tailor_calls"].call_count == 2


def test_tailor_docx_resume_updated_since_last_tailor_regenerates_this_jobs_files(env):
    _run(_tailor_docx_argv(env))
    _, cover_letter_path = _tailored_docx_paths("Acme AI", env["job_id"], str(env["out_dir"]))
    old_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text

    env["tailor_calls"].return_value = _fake_tailored_application("_V2")
    _add_profile_version(env["profile_db"], RESUME_TEXT + " Updated with a new AWS certification.")

    _run(_tailor_docx_argv(env))

    new_cover_text = docx.Document(str(cover_letter_path)).paragraphs[0].text
    assert new_cover_text != old_cover_text
    assert env["tailor_calls"].call_count == 2
    assert env["paragraph_calls"].call_count == 2

    row = _get_job_row(env["jobs_db"], env["job_id"])
    assert row["tailor_evidence_notes"] == '["evidence note_V2"]'


def test_tailor_docx_files_deleted_manually_is_treated_as_a_cache_miss(env):
    _run(_tailor_docx_argv(env))
    resume_path, _ = _tailored_docx_paths("Acme AI", env["job_id"], str(env["out_dir"]))
    resume_path.unlink()

    _run(_tailor_docx_argv(env))

    assert resume_path.exists()
    assert env["tailor_calls"].call_count == 2


def test_tailor_docx_force_flag_regenerates_even_when_hash_matches_and_files_exist(env):
    _run(_tailor_docx_argv(env))
    _run(_tailor_docx_argv(env, force=True))

    assert env["tailor_calls"].call_count == 2
    assert env["paragraph_calls"].call_count == 2


def test_tailor_docx_rejects_a_concurrent_call_while_the_lock_is_held(env):
    # Simulates another session (e.g. a second browser tab) already mid-run
    # for the same job - the DB-backed lock is the only thing that can catch
    # this, since each `_run()` here opens and closes its own connection.
    conn = connect_jobs(env["jobs_db"])
    try:
        assert try_claim_tailoring_lock(conn, env["job_id"]) is not None
    finally:
        conn.close()

    with pytest.raises(SystemExit, match="already in progress"):
        _run(_tailor_docx_argv(env))

    assert env["tailor_calls"].call_count == 0
    assert env["paragraph_calls"].call_count == 0


def test_tailor_docx_releases_the_lock_even_when_generation_raises_so_a_retry_can_claim_it(env):
    env["paragraph_calls"].side_effect = RuntimeError("simulated LLM failure")

    with pytest.raises(RuntimeError):
        _run(_tailor_docx_argv(env))

    row = _get_job_row(env["jobs_db"], env["job_id"])
    assert row["tailoring_lock_started_at"] is None

    env["paragraph_calls"].side_effect = None  # simulate the retry succeeding
    _run(_tailor_docx_argv(env))

    assert env["tailor_calls"].call_count == 2


def test_tailor_docx_releases_the_lock_when_the_earlier_llm_call_raises(env):
    # A distinct failure site from the test above - `generate_tailored_application`
    # raises before `generate_paragraph_edits` is ever reached, exercising the
    # same `finally` release from a different point in the try block, so a
    # regression that only breaks release on this path (e.g. an accidental
    # early `return` added between the two LLM calls) would still be caught.
    env["tailor_calls"].side_effect = RuntimeError("simulated LLM failure")

    with pytest.raises(RuntimeError):
        _run(_tailor_docx_argv(env))

    row = _get_job_row(env["jobs_db"], env["job_id"])
    assert row["tailoring_lock_started_at"] is None
    assert env["paragraph_calls"].call_count == 0

    env["tailor_calls"].side_effect = None
    _run(_tailor_docx_argv(env))

    assert env["tailor_calls"].call_count == 2


def test_tailor_docx_unknown_job_id_raises_system_exit(env):
    """Driven through the real CLI wiring (`build_parser()`/`args.func`),
    matching this file's own stated test philosophy - not just the
    internal `_tailor_docx_for_job` helper directly."""
    unknown_job_id = env["job_id"] + 999
    with pytest.raises(SystemExit, match=f"No job #{unknown_job_id} found"):
        _run(_tailor_docx_argv(env, job_id=unknown_job_id))


def test_tailor_docx_cache_hit_succeeds_even_after_source_docx_is_deleted(env, monkeypatch):
    """The frozen cache-hit contract requires no source-.docx read at all on
    a hit. Delete the source file the first generation read from, then
    monkeypatch `_find_source_resume_docx` to fail loudly if called - the
    second run must succeed anyway (no SystemExit, no crash) and return the
    exact same warning that was stored at generation time."""
    _run(_tailor_docx_argv(env))
    stored_warning_before = _get_job_row(env["jobs_db"], env["job_id"])["tailor_page_risk_warning"]

    (env["resume_dir"] / "resume.docx").unlink()

    def _fail_if_called(*args, **kwargs):
        raise AssertionError("cache hit must not read the source .docx")

    monkeypatch.setattr(jobs_cli, "_find_source_resume_docx", _fail_if_called)

    _run(_tailor_docx_argv(env))  # must not raise SystemExit or AssertionError

    stored_warning_after = _get_job_row(env["jobs_db"], env["job_id"])["tailor_page_risk_warning"]
    assert stored_warning_after == stored_warning_before
    assert env["tailor_calls"].call_count == 1


# --------------------------------------------------------------------------
# `tailor` - plain-text, deliberately uncached
# --------------------------------------------------------------------------


def test_cmd_tailor_calls_llm_both_times_and_overwrites_txt_output(env, tmp_path):
    txt_out_dir = tmp_path / "tailored_txt"
    argv = [
        "tailor",
        str(env["job_id"]),
        "--db",
        str(env["jobs_db"]),
        "--profile-db",
        str(env["profile_db"]),
        "--out-dir",
        str(txt_out_dir),
    ]

    _run(argv)
    _run(argv)

    assert env["tailor_calls"].call_count == 2
    resume_txt = txt_out_dir / f"{env['job_id']}_resume.txt"
    assert resume_txt.read_text(encoding="utf-8") == "TAILORED RESUME"

    # The plain-text path must never touch the docx-cache's DB fields -
    # sharing that state is exactly the bug this refactor's amendment fixed.
    row = _get_job_row(env["jobs_db"], env["job_id"])
    assert row["tailor_hash"] is None


def test_cmd_tailor_unknown_job_id_raises_system_exit(env, tmp_path):
    """Driven through the real CLI wiring (`build_parser()`/`args.func`),
    matching this file's own stated test philosophy - not just the
    internal helper functions directly."""
    unknown_job_id = env["job_id"] + 999
    argv = [
        "tailor",
        str(unknown_job_id),
        "--db",
        str(env["jobs_db"]),
        "--profile-db",
        str(env["profile_db"]),
        "--out-dir",
        str(tmp_path / "tailored_txt"),
    ]

    with pytest.raises(SystemExit, match=f"No job #{unknown_job_id} found"):
        _run(argv)


# --------------------------------------------------------------------------
# `migrate-legacy-tailoring`
# --------------------------------------------------------------------------


def test_migrate_legacy_tailoring_writes_txt_from_legacy_db_text_and_is_idempotent(tmp_path):
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name="Bending Spoons")
    _add_legacy_columns_with_text(jobs_db, job_id, "LEGACY RESUME TEXT", "LEGACY COVER LETTER TEXT")

    out_dir = tmp_path / "tailored_txt"
    generated_cv_dir = tmp_path / "generated_cv"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    resume_txt = out_dir / f"{job_id}_resume.txt"
    cover_txt = out_dir / f"{job_id}_cover_letter.txt"
    assert resume_txt.read_text(encoding="utf-8") == "LEGACY RESUME TEXT"
    assert cover_txt.read_text(encoding="utf-8") == "LEGACY COVER LETTER TEXT"

    _run(argv)  # idempotent - no error, content unchanged
    assert resume_txt.read_text(encoding="utf-8") == "LEGACY RESUME TEXT"
    assert cover_txt.read_text(encoding="utf-8") == "LEGACY COVER LETTER TEXT"


def test_migrate_legacy_tailoring_partial_state_only_writes_the_missing_txt_file(tmp_path):
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name="Acme AI")
    _add_legacy_columns_with_text(jobs_db, job_id, "LEGACY RESUME TEXT", "LEGACY COVER LETTER TEXT")

    out_dir = tmp_path / "tailored_txt"
    out_dir.mkdir()
    # Simulate the resume half of the .txt export having already happened
    # (e.g. from an earlier partial run) - the migration must not clobber it.
    (out_dir / f"{job_id}_resume.txt").write_text("ALREADY EXPORTED RESUME", encoding="utf-8")

    generated_cv_dir = tmp_path / "generated_cv"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    assert (out_dir / f"{job_id}_resume.txt").read_text(encoding="utf-8") == "ALREADY EXPORTED RESUME"
    assert (out_dir / f"{job_id}_cover_letter.txt").read_text(encoding="utf-8") == "LEGACY COVER LETTER TEXT"


def test_migrate_legacy_tailoring_renames_unambiguous_legacy_docx(tmp_path):
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name="Bending Spoons")

    generated_cv_dir = tmp_path / "generated_cv"
    company_dir = generated_cv_dir / _sanitize_filename("Bending Spoons")
    company_dir.mkdir(parents=True)
    (company_dir / "resume.docx").write_bytes(b"old resume bytes")
    (company_dir / "cover_letter.docx").write_bytes(b"old cover bytes")

    out_dir = tmp_path / "tailored_txt"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    assert (company_dir / f"{job_id}_resume.docx").read_bytes() == b"old resume bytes"
    assert (company_dir / f"{job_id}_cover_letter.docx").read_bytes() == b"old cover bytes"
    assert not (company_dir / "resume.docx").exists()
    assert not (company_dir / "cover_letter.docx").exists()

    _run(argv)  # idempotent - old-format files are already gone, no error


def test_migrate_legacy_tailoring_ambiguous_docx_folder_warns_and_leaves_files_untouched(tmp_path, capsys):
    jobs_db = tmp_path / "jobs.db"
    _make_job(jobs_db, company_name="Bending Spoons")
    _make_job(jobs_db, company_name="Bending Spoons")  # two jobs -> ambiguous match

    generated_cv_dir = tmp_path / "generated_cv"
    company_dir = generated_cv_dir / _sanitize_filename("Bending Spoons")
    company_dir.mkdir(parents=True)
    (company_dir / "resume.docx").write_bytes(b"old resume bytes")
    (company_dir / "cover_letter.docx").write_bytes(b"old cover bytes")

    out_dir = tmp_path / "tailored_txt"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    captured = capsys.readouterr()
    assert "ambiguous" in captured.out.lower()
    assert (company_dir / "resume.docx").exists()
    assert (company_dir / "cover_letter.docx").exists()


def test_migrate_legacy_tailoring_zero_matching_jobs_warns_and_leaves_files_untouched(tmp_path, capsys):
    jobs_db = tmp_path / "jobs.db"
    _make_job(jobs_db, company_name="Some Other Company")  # no job matches "Bending Spoons"

    generated_cv_dir = tmp_path / "generated_cv"
    company_dir = generated_cv_dir / _sanitize_filename("Bending Spoons")
    company_dir.mkdir(parents=True)
    (company_dir / "resume.docx").write_bytes(b"old resume bytes")

    out_dir = tmp_path / "tailored_txt"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    captured = capsys.readouterr()
    # Zero matches is orphaned data with nothing to disambiguate - distinct
    # wording from the genuinely-ambiguous (2+ match) case above.
    assert "no matching job" in captured.out.lower()
    assert "ambiguous" not in captured.out.lower()
    assert (company_dir / "resume.docx").exists()


def test_migrate_legacy_tailoring_matches_company_less_job_via_fallback_slug(tmp_path):
    """A legacy folder created under the `job_{id}`/`unknown_company`
    fallback slug (for a job with no company_name) must still be matched
    and renamed like any other unambiguous case - not permanently reported
    ambiguous just because `company_name` is falsy."""
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name=None)

    generated_cv_dir = tmp_path / "generated_cv"
    company_dir = generated_cv_dir / f"job_{job_id}"
    company_dir.mkdir(parents=True)
    (company_dir / "resume.docx").write_bytes(b"old resume bytes")
    (company_dir / "cover_letter.docx").write_bytes(b"old cover bytes")

    out_dir = tmp_path / "tailored_txt"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)

    assert (company_dir / f"{job_id}_resume.docx").exists()
    assert (company_dir / f"{job_id}_cover_letter.docx").exists()
    assert not (company_dir / "resume.docx").exists()


def test_migrate_legacy_tailoring_continues_after_oserror_renaming_one_folder(tmp_path, monkeypatch):
    """One OSError (locked file, permission denied) on one company folder
    must not abort the entire migration loop for every subsequent folder."""
    jobs_db = tmp_path / "jobs.db"
    broken_job_id = _make_job(jobs_db, company_name="Broken Co")
    good_job_id = _make_job(jobs_db, company_name="Good Co")

    generated_cv_dir = tmp_path / "generated_cv"
    broken_dir = generated_cv_dir / _sanitize_filename("Broken Co")
    good_dir = generated_cv_dir / _sanitize_filename("Good Co")
    broken_dir.mkdir(parents=True)
    good_dir.mkdir(parents=True)
    (broken_dir / "resume.docx").write_bytes(b"broken resume")
    (broken_dir / "cover_letter.docx").write_bytes(b"broken cover")
    (good_dir / "resume.docx").write_bytes(b"good resume")
    (good_dir / "cover_letter.docx").write_bytes(b"good cover")

    original_rename = Path.rename

    def flaky_rename(self, target):
        if self.parent.name == _sanitize_filename("Broken Co"):
            raise OSError("simulated permission denied")
        return original_rename(self, target)

    monkeypatch.setattr(Path, "rename", flaky_rename)

    out_dir = tmp_path / "tailored_txt"
    argv = [
        "migrate-legacy-tailoring",
        "--db",
        str(jobs_db),
        "--out-dir",
        str(out_dir),
        "--generated-cv-dir",
        str(generated_cv_dir),
    ]

    _run(argv)  # must not raise - and must still process Good_Co

    assert (broken_dir / "resume.docx").exists()
    assert not (broken_dir / f"{broken_job_id}_resume.docx").exists()

    assert (good_dir / f"{good_job_id}_resume.docx").exists()
    assert (good_dir / f"{good_job_id}_cover_letter.docx").exists()


# --------------------------------------------------------------------------
# `outreach` - drafted message text is written to a file, not stored in the DB
# --------------------------------------------------------------------------


def _add_narrative_and_profile(profile_db_path: Path) -> None:
    _add_profile_version(profile_db_path, RESUME_TEXT)
    conn = connect_profile(profile_db_path)
    try:
        insert_narrative(conn, "Why AI, why UK, why them - the candidate's narrative core.")
    finally:
        conn.close()


@pytest.fixture
def outreach_env(tmp_path, monkeypatch):
    """A ready-to-use jobs.db (one job with a recruiter) + profile.db (resume
    + narrative core), with `draft_outreach_message` monkeypatched so tests
    can assert on the file it writes without hitting a real Gemini API."""
    jobs_db = tmp_path / "jobs.db"
    profile_db = tmp_path / "profile.db"
    out_dir = tmp_path / "generated_cv"

    conn = connect_jobs(jobs_db)
    try:
        job_id = insert_job(
            conn,
            "job posting text",
            JobExtraction(
                job_title="AI Engineer",
                company_name="Bending Spoons",
                is_agency_posting=False,
                recruiter_name="Sarah Cole",
                recruiter_contact="sarah@bendingspoons.com",
            ),
        )
    finally:
        conn.close()

    _add_narrative_and_profile(profile_db)

    draft_calls = MagicMock(return_value=OutreachDraft(message="Hi Sarah, I'd love to chat about the AI Engineer role."))
    monkeypatch.setattr(jobs_cli, "draft_outreach_message", draft_calls)

    return {
        "jobs_db": jobs_db,
        "profile_db": profile_db,
        "out_dir": out_dir,
        "job_id": job_id,
        "draft_calls": draft_calls,
    }


def test_cmd_outreach_writes_the_expected_txt_file_and_stores_no_message_column(outreach_env):
    argv = [
        "outreach",
        str(outreach_env["job_id"]),
        "--channel",
        LINKEDIN_NOTE,
        "--db",
        str(outreach_env["jobs_db"]),
        "--profile-db",
        str(outreach_env["profile_db"]),
        "--out-dir",
        str(outreach_env["out_dir"]),
    ]

    _run(argv)

    conn = connect_jobs(outreach_env["jobs_db"])
    try:
        ensure_outreach_schema(conn)
        messages = list_outreach_messages(conn, outreach_env["job_id"])
    finally:
        conn.close()

    assert len(messages) == 1
    message_row = messages[0]
    assert "message" not in message_row.keys()
    assert message_row["char_count"] == len("Hi Sarah, I'd love to chat about the AI Engineer role.")

    path = _outreach_message_path(
        "Bending Spoons", outreach_env["job_id"], LINKEDIN_NOTE, message_row["id"], str(outreach_env["out_dir"])
    )
    assert path.exists()
    assert path.read_text(encoding="utf-8") == "Hi Sarah, I'd love to chat about the AI Engineer role."


def test_cmd_outreach_second_draft_to_same_job_and_channel_gets_its_own_file(outreach_env):
    """`outreach_messages` is an insert-only history table - a second draft
    to the same job+channel (e.g. to a different contact) must not
    overwrite the first draft's file."""
    argv = [
        "outreach",
        str(outreach_env["job_id"]),
        "--channel",
        LINKEDIN_NOTE,
        "--db",
        str(outreach_env["jobs_db"]),
        "--profile-db",
        str(outreach_env["profile_db"]),
        "--out-dir",
        str(outreach_env["out_dir"]),
    ]

    _run(argv)
    outreach_env["draft_calls"].return_value = OutreachDraft(message="A completely different second draft.")
    _run(argv)

    conn = connect_jobs(outreach_env["jobs_db"])
    try:
        ensure_outreach_schema(conn)
        messages = list_outreach_messages(conn, outreach_env["job_id"])
    finally:
        conn.close()

    assert len(messages) == 2
    paths = [
        _outreach_message_path("Bending Spoons", outreach_env["job_id"], LINKEDIN_NOTE, row["id"], str(outreach_env["out_dir"]))
        for row in messages
    ]
    assert paths[0] != paths[1]
    assert all(p.exists() for p in paths)


# --------------------------------------------------------------------------
# `migrate-legacy-outreach`
# --------------------------------------------------------------------------


def _add_legacy_outreach_message_column_with_text(jobs_db_path: Path, job_id: int, channel: str, message: str) -> int:
    """Simulate a jobs.db created before this refactor: `outreach_messages`
    still has a real `message` column with real data (only a one-time
    migration drops it - schema application never does)."""
    conn = sqlite3.connect(jobs_db_path)
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS outreach_messages (
            id INTEGER PRIMARY KEY,
            job_id INTEGER NOT NULL,
            contact_id INTEGER,
            contact_name TEXT NOT NULL,
            channel TEXT NOT NULL,
            message TEXT NOT NULL,
            char_count INTEGER NOT NULL,
            created_at TEXT NOT NULL
        )
        """
    )
    cursor = conn.execute(
        """
        INSERT INTO outreach_messages (job_id, contact_id, contact_name, channel, message, char_count, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (job_id, None, "Sarah Cole", channel, message, len(message), "2026-01-01T00:00:00+00:00"),
    )
    conn.commit()
    message_id = cursor.lastrowid
    conn.close()
    return message_id


def test_migrate_legacy_outreach_writes_txt_from_legacy_db_text_and_is_idempotent(tmp_path):
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name="Bending Spoons")
    message_id = _add_legacy_outreach_message_column_with_text(jobs_db, job_id, LINKEDIN_NOTE, "LEGACY MESSAGE TEXT")

    out_dir = tmp_path / "generated_cv"
    argv = ["migrate-legacy-outreach", "--db", str(jobs_db), "--out-dir", str(out_dir)]

    _run(argv)

    path = _outreach_message_path("Bending Spoons", job_id, LINKEDIN_NOTE, message_id, str(out_dir))
    assert path.read_text(encoding="utf-8") == "LEGACY MESSAGE TEXT"

    conn = sqlite3.connect(jobs_db)
    existing = {row[1] for row in conn.execute("PRAGMA table_info(outreach_messages)")}
    conn.close()
    assert "message" not in existing

    _run(argv)  # idempotent - column already dropped, file already written, no error
    assert path.read_text(encoding="utf-8") == "LEGACY MESSAGE TEXT"


def test_migrate_legacy_outreach_on_a_fresh_db_is_a_no_op(tmp_path, capsys):
    jobs_db = tmp_path / "jobs.db"
    _make_job(jobs_db, company_name="Acme AI")

    out_dir = tmp_path / "generated_cv"
    argv = ["migrate-legacy-outreach", "--db", str(jobs_db), "--out-dir", str(out_dir)]

    _run(argv)  # must not raise

    captured = capsys.readouterr()
    assert "nothing to do" in captured.out.lower()


def test_migrate_legacy_outreach_then_new_insert_succeeds(tmp_path):
    """The whole point of dropping the column: a subsequent draft insert
    must succeed against the migrated table (a `message NOT NULL` legacy
    column left in place would break inserts that no longer supply it)."""
    jobs_db = tmp_path / "jobs.db"
    job_id = _make_job(jobs_db, company_name="Bending Spoons")
    _add_legacy_outreach_message_column_with_text(jobs_db, job_id, LINKEDIN_NOTE, "LEGACY MESSAGE TEXT")

    out_dir = tmp_path / "generated_cv"
    _run(["migrate-legacy-outreach", "--db", str(jobs_db), "--out-dir", str(out_dir)])

    from jobs.outreach_db import insert_outreach_message

    conn = connect_jobs(jobs_db)
    try:
        ensure_outreach_schema(conn)
        new_id = insert_outreach_message(
            conn, job_id, contact_id=None, contact_name="New Contact", channel=EMAIL, message="A brand new message."
        )
        assert new_id is not None
    finally:
        conn.close()


def test_drafting_against_an_unmigrated_legacy_db_raises_a_friendly_error_naming_the_migration_command(
    tmp_path, monkeypatch
):
    """Without `migrate-legacy-outreach` having been run, the pre-existing
    `message NOT NULL` column would otherwise surface as a raw, unactionable
    `sqlite3.IntegrityError` - this must instead point the user at the fix."""
    jobs_db = tmp_path / "jobs.db"
    profile_db = tmp_path / "profile.db"
    conn = connect_jobs(jobs_db)
    try:
        job_id = insert_job(
            conn,
            "job posting text",
            JobExtraction(
                job_title="AI Engineer",
                company_name="Bending Spoons",
                is_agency_posting=False,
                recruiter_name="Sarah Cole",
                recruiter_contact="sarah@bendingspoons.com",
            ),
        )
    finally:
        conn.close()
    _add_legacy_outreach_message_column_with_text(jobs_db, job_id, LINKEDIN_NOTE, "LEGACY MESSAGE TEXT")
    _add_narrative_and_profile(profile_db)

    monkeypatch.setattr(
        jobs_cli, "draft_outreach_message", MagicMock(return_value=OutreachDraft(message="A brand new draft."))
    )

    with pytest.raises(SystemExit, match="migrate-legacy-outreach"):
        _run(
            [
                "outreach",
                str(job_id),
                "--channel",
                LINKEDIN_NOTE,
                "--db",
                str(jobs_db),
                "--profile-db",
                str(profile_db),
                "--out-dir",
                str(tmp_path / "generated_cv"),
            ]
        )


def test_atomic_write_text_writes_full_content_and_leaves_no_tmp_file(tmp_path):
    target = tmp_path / "file.txt"

    jobs_cli._atomic_write_text(target, "hello world")

    assert target.read_text(encoding="utf-8") == "hello world"
    assert list(tmp_path.glob("*.tmp")) == []


def test_atomic_write_text_leaves_prior_content_intact_and_cleans_up_tmp_when_replace_fails(tmp_path, monkeypatch):
    target = tmp_path / "file.txt"
    target.write_text("original content", encoding="utf-8")

    monkeypatch.setattr(jobs_cli.os, "replace", MagicMock(side_effect=OSError("simulated disk failure")))

    with pytest.raises(OSError, match="simulated disk failure"):
        jobs_cli._atomic_write_text(target, "new content that must never land")

    assert target.read_text(encoding="utf-8") == "original content"
    assert list(tmp_path.glob("*.tmp")) == []


def test_atomic_write_text_leaves_prior_content_intact_and_cleans_up_tmp_when_fsync_fails(tmp_path, monkeypatch):
    """A real ENOSPC realistically surfaces during the write/fsync step, not
    `os.replace` - covered separately from the replace-failure test above so
    both halves of the helper's error path are exercised."""
    target = tmp_path / "file.txt"
    target.write_text("original content", encoding="utf-8")

    monkeypatch.setattr(jobs_cli.os, "fsync", MagicMock(side_effect=OSError("simulated ENOSPC")))

    with pytest.raises(OSError, match="simulated ENOSPC"):
        jobs_cli._atomic_write_text(target, "new content that must never land")

    assert target.read_text(encoding="utf-8") == "original content"
    assert list(tmp_path.glob("*.tmp")) == []


def test_atomic_write_text_successive_calls_use_distinct_tmp_filenames(tmp_path, monkeypatch):
    """Each call must get its own pid+uuid-suffixed tmp sibling, not a
    single static name - otherwise two overlapping writers targeting the
    same path would share one tmp file and could interleave/corrupt it."""
    target = tmp_path / "file.txt"
    seen_tmp_names = []
    real_open = builtins.open

    def spying_open(path, *args, **kwargs):
        if Path(path) != target:
            seen_tmp_names.append(str(path))
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr("builtins.open", spying_open)

    jobs_cli._atomic_write_text(target, "first")
    jobs_cli._atomic_write_text(target, "second")

    assert len(seen_tmp_names) == 2
    assert seen_tmp_names[0] != seen_tmp_names[1]
    assert target.read_text(encoding="utf-8") == "second"


def test_atomic_write_text_fsyncs_the_directory_only_after_a_successful_replace(tmp_path, monkeypatch):
    # jobs.atomic_fs._fsync_directory itself is unit-tested in
    # test_atomic_fs.py - this only proves _atomic_write_text calls it with
    # the right directory, and only once the rename has actually succeeded.
    target = tmp_path / "file.txt"
    mock_fsync_directory = MagicMock()
    monkeypatch.setattr(jobs_cli, "_fsync_directory", mock_fsync_directory)

    jobs_cli._atomic_write_text(target, "hello world")

    mock_fsync_directory.assert_called_once_with(tmp_path)
    assert target.read_text(encoding="utf-8") == "hello world"


def test_atomic_write_text_does_not_fsync_the_directory_when_replace_fails(tmp_path, monkeypatch):
    target = tmp_path / "file.txt"
    mock_fsync_directory = MagicMock()
    monkeypatch.setattr(jobs_cli, "_fsync_directory", mock_fsync_directory)
    monkeypatch.setattr(jobs_cli.os, "replace", MagicMock(side_effect=OSError("simulated disk failure")))

    with pytest.raises(OSError, match="simulated disk failure"):
        jobs_cli._atomic_write_text(target, "hello world")

    mock_fsync_directory.assert_not_called()


def test_cmd_outreach_write_failure_after_db_commit_raises_distinct_lost_text_error(outreach_env, monkeypatch, capsys):
    """The DB row commits before the file write is attempted (see
    `_draft_and_store_outreach`), so a write failure here is worse than an
    ordinary I/O error - the metadata now refers to text that no longer
    exists anywhere. The error must say so explicitly, not just surface a
    generic OSError, and the drafted text must be printed for manual
    recovery since re-running would call the LLM again from scratch."""
    monkeypatch.setattr(jobs_cli, "_atomic_write_text", MagicMock(side_effect=OSError("disk full")))

    with pytest.raises(SystemExit, match="was logged to the database.*text itself was not saved"):
        _run(
            [
                "outreach",
                str(outreach_env["job_id"]),
                "--channel",
                LINKEDIN_NOTE,
                "--db",
                str(outreach_env["jobs_db"]),
                "--profile-db",
                str(outreach_env["profile_db"]),
                "--out-dir",
                str(outreach_env["out_dir"]),
            ]
        )

    conn = connect_jobs(outreach_env["jobs_db"])
    try:
        row = list_outreach_messages(conn, outreach_env["job_id"])[0]
        assert row["write_failed_at"] is not None  # marks it as distinguishable from a later-deleted file
    finally:
        conn.close()

    assert "Hi Sarah, I'd love to chat about the AI Engineer role." in capsys.readouterr().out

    # The DB row is still there (insert-only convention, no rollback) -
    # this is the exact orphaned-metadata state the error message must warn about.
    conn = connect_jobs(outreach_env["jobs_db"])
    try:
        ensure_outreach_schema(conn)
        messages = list_outreach_messages(conn, outreach_env["job_id"])
    finally:
        conn.close()
    assert len(messages) == 1


def test_cmd_outreach_write_failure_survives_the_write_failure_marker_itself_failing(outreach_env, monkeypatch):
    # mark_outreach_write_failed runs after the file write already failed -
    # if it also raises (e.g. the same disk backs this DB), the original,
    # more informative SystemExit (with the recovery text) must still be
    # what the operator sees, not a raw sqlite error replacing it.
    monkeypatch.setattr(jobs_cli, "_atomic_write_text", MagicMock(side_effect=OSError("disk full")))
    monkeypatch.setattr(jobs_cli, "mark_outreach_write_failed", MagicMock(side_effect=sqlite3.OperationalError("disk full")))

    with pytest.raises(SystemExit, match="was logged to the database.*text itself was not saved"):
        _run(
            [
                "outreach",
                str(outreach_env["job_id"]),
                "--channel",
                LINKEDIN_NOTE,
                "--db",
                str(outreach_env["jobs_db"]),
                "--profile-db",
                str(outreach_env["profile_db"]),
                "--out-dir",
                str(outreach_env["out_dir"]),
            ]
        )
